"""Offer 捕手 - Web 控制台"""
from flask import Flask, render_template, request, jsonify, send_file
from flask_cors import CORS
from engine.matcher import MatchEngine
import os, json, re, requests

app = Flask(__name__)
CORS(app)

# ===== API 配置（运行时可变，用户自行填写 Key） =====
api_config = {
    "text_provider": "deepseek",
    "text_key": "",
    "text_url": "https://api.deepseek.com/v1/chat/completions",
    "text_model": "deepseek-chat",
    "vision_provider": "dashscope",
    "vision_key": "",
    "vision_url": "https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions",
    "vision_model": "qwen-vl-plus",
}

def text_api_url():  return api_config["text_url"]
def text_api_key():  return api_config["text_key"]
def text_model():    return api_config["text_model"]
def vision_api_url(): return api_config["vision_url"]
def vision_api_key(): return api_config["vision_key"]
def vision_model():  return api_config["vision_model"]


# 内存存储（Demo 用）
state = {
    "resume": None, "resume_text": "", "jds": [],
    "weight": {"exp": 40, "hard": 25, "skill": 15, "company": 10, "fit": 10},
    "config": {"ds_key": "", "dashscope_key": ""}  # 用户通过网页设置的 Key
}

engine = None

def get_ds_key():
    """优先使用网页配置的 Key，否则尝试环境变量"""
    return state["config"].get("ds_key") or os.environ.get("DS_KEY", "")

def get_dashscope_key():
    """优先使用网页配置的 Key，否则尝试环境变量"""
    return state["config"].get("dashscope_key") or os.environ.get("DASHSCOPE_KEY", "")

def get_engine():
    """懒加载 MatchEngine，Key 变化时重建"""
    global engine
    key = get_ds_key()
    if not key:
        return None
    if engine is None or engine.key != key:
        engine = MatchEngine(key)
    return engine

def require_engine():
    """需要 DeepSeek Key 的操作前校验"""
    eng = get_engine()
    if not eng:
        raise ValueError("请先在左侧面板配置 DeepSeek API Key")
    return eng

def require_vision():
    """需要视觉 API Key 的操作前校验（支持 DashScope/HunYuan/自定义）"""
    key = vision_api_key()
    if not key:
        raise ValueError("请先在 API 配置面板填写视觉 API Key")
    return key

@app.route("/")
def index():
    return render_template("index.html")

# ===== API 配置（用户通过网页设置自己的 Key 和提供商） =====
@app.route("/api/config", methods=["GET", "POST"])
def config_handler():
    if request.method == "POST":
        data = request.json or {}
        # 旧格式兼容
        if "ds_key" in data:
            state["config"]["ds_key"] = data["ds_key"].strip()
        if "dashscope_key" in data:
            state["config"]["dashscope_key"] = data["dashscope_key"].strip()
        # 新格式：支持切换提供商
        for k in ["text_key", "text_url", "text_model",
                  "vision_key", "vision_url", "vision_model"]:
            if k in data:
                api_config[k] = data[k].strip() if data[k] else data[k]
        # 同步到 state.config
        if api_config.get("text_key"):
            state["config"]["ds_key"] = api_config["text_key"]
        if api_config.get("vision_key"):
            state["config"]["dashscope_key"] = api_config["vision_key"]
        global engine; engine = None
        return jsonify({"ok": True,
                        "ds_configured": bool(get_ds_key()),
                        "dashscope_configured": bool(vision_api_key())})
    # GET: 返回当前配置（不暴露完整 Key）
    return jsonify({
        "ds_configured": bool(get_ds_key()),
        "dashscope_configured": bool(vision_api_key()),
        "text_provider": api_config["text_provider"],
        "text_model": api_config["text_model"],
        "vision_model": api_config["vision_model"],
        "text_key_preview": api_config.get("text_key","")[:8] + "***" if api_config.get("text_key") else "",
        "vision_key_preview": api_config.get("vision_key","")[:8] + "***" if api_config.get("vision_key") else "",
    })

# ===== 简历 =====
@app.route("/api/resume/parse", methods=["POST"])
def parse_resume():
    try: eng = require_engine()
    except ValueError as e: return jsonify({"error": str(e)}), 400
    data = request.json or {}
    text = data.get("text", "").strip()
    if not text:
        return jsonify({"error": "简历为空"}), 400
    state["resume"] = eng.parse_resume(text)
    state["resume_text"] = text
    return jsonify({"ok": True, "resume": state["resume"]})

# ===== 简历 VL 直接解析（客户端渲染图片 → VL 提取结构化数据） =====
@app.route("/api/resume/vision_parse", methods=["POST"])
def resume_vision_parse():
    data = request.json or {}
    images = data.get("images", [])
    if not images:
        return jsonify({"error": "图片数据为空"}), 400

    # 去掉 data URL 前缀
    cleaned = []
    for img in images:
        if isinstance(img, str) and "," in img:
            img = img.split(",", 1)[1]
        cleaned.append(img)

    parsed = call_vl_for_resume(cleaned)
    if not parsed or not parsed.get("name"):
        return jsonify({"error": "VL 未能识别简历信息，请确认图片清晰可读"}), 500

    state["resume"] = parsed
    state["resume_text"] = json.dumps(parsed, ensure_ascii=False)
    return jsonify({"ok": True, "resume": parsed, "source": "VL直出"})


# ===== 简历视觉清洗（网站粘贴乱码 → 图片 → VL 识别） =====
@app.route("/api/resume/vision", methods=["POST"])
def resume_vision():
    data = request.json or {}
    img_b64 = data.get("image", "")
    if not img_b64:
        return jsonify({"error": "图片数据为空"}), 400

    if "," in img_b64:
        img_b64 = img_b64.split(",", 1)[1]

    try:
        api_key = require_vision()
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    try:
        resp = requests.post(
            vision_api_url(),
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            json={
                "model": vision_model(),
                "messages": [{"role": "user", "content": [
                    {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_b64}"}},
                    {"type": "text", "text": (
                        "这是一份从网站粘贴出来的简历文本截图，其中混入了大量乱码字符（如随机字母数字串、"
                        "单字符碎片、hash 值等水印）。请提取其中的真实简历内容，去掉所有乱码，"
                        "保持原有的段落结构和换行。只返回清洗后的纯文本，不要加任何解释。"
                    )},
                ]}],
            },
            timeout=30,
        )
        if resp.status_code != 200:
            return jsonify({"error": f"VL 返回 {resp.status_code}"}), 500

        result = resp.json()
        text = result.get("choices", [{}])[0].get("message", {}).get("content", "")
        text = text.strip()
        # 去掉可能的 markdown 代码块包裹
        if text.startswith("```"):
            text = text.split("\n", 1)[-1]
            if text.endswith("```"): text = text[:-3]
        text = text.strip()

        if len(text) < 20:
            return jsonify({"error": "VL 提取结果太短"}), 500

        return jsonify({"ok": True, "text": text})
    except Exception as e:
        return jsonify({"error": f"VL 调用异常: {str(e)}"}), 500


# ===== 简历解析：文本优先（DeepSeek）+ VL 兜底 =====
@app.route("/api/resume/upload", methods=["POST"])
def upload_resume():
    if "file" not in request.files:
        return jsonify({"error": "未选择文件"}), 400
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "文件名为空"}), 400

    ext = os.path.splitext(file.filename)[1].lower()
    file_bytes = file.read()
    import base64

    try:
        text = ""
        parsed = None
        source = ""

        if ext == ".pdf":
            # PDF → pdfplumber 提取文字 → DeepSeek 解析
            import pdfplumber, io as _io4
            with pdfplumber.open(_io4.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t: text += t + "\n"
            if len(text.strip()) > 200:
                parsed = call_ds_for_resume(text)
                source = "PDF文字→DeepSeek"
            else:
                # 扫描版 PDF，让客户端用 PDF.js 渲染
                return jsonify({"needs_vision": True, "text": text.strip()[:100]}), 200

        elif ext == ".docx":
            from docx import Document
            import io as _io5
            doc = Document(_io5.BytesIO(file_bytes))
            # 提取段落文字
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            # 提取表格文字
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text += "\n" + row_text
            # 还不够？深挖 XML 里的所有文字（包括文本框、页眉页脚）
            if len(text.strip()) < 100:
                import zipfile, xml.etree.ElementTree as ET
                with zipfile.ZipFile(_io5.BytesIO(file_bytes)) as z:
                    if "word/document.xml" in z.namelist():
                        tree = ET.parse(z.open("word/document.xml"))
                        for wt in tree.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
                            if wt.text and wt.text.strip():
                                text += wt.text.strip() + "\n"
            if len(text.strip()) > 50:
                parsed = call_ds_for_resume(text)
                source = "Word→DeepSeek"
            else:
                return jsonify({"error": "Word 文件无法提取内容（可能全为文本框），请导出为 PDF 后上传"}), 400

        elif ext == ".txt":
            text = file_bytes.decode("utf-8", errors="ignore")
            if len(text.strip()) > 50:
                parsed = call_ds_for_resume(text)
                source = "TXT→DeepSeek"
            else:
                return jsonify({"error": "文本文件内容为空"}), 400

        elif ext in (".png", ".jpg", ".jpeg"):
            # 图片直接走 VL
            images_b64 = [base64.b64encode(file_bytes).decode()]
            parsed = call_vl_for_resume(images_b64)
            source = "图片→VL"

        else:
            return jsonify({"error": f"不支持的文件格式: {ext}"}), 400

        # 名字可能因 WPS 文本框提取不到，有教育或经历就算成功
        has_content = parsed and (parsed.get("name") or parsed.get("education") or parsed.get("experience"))
        if has_content:
            state["resume"] = parsed
            state["resume_text"] = json.dumps(parsed, ensure_ascii=False)
            return jsonify({"ok": True, "resume": parsed, "source": source})
        else:
            return jsonify({"error": "未能识别简历信息，请确认文件内容完整"}), 500

    except Exception as e:
        return jsonify({"error": f"文件解析失败: {str(e)}"}), 500


# ===== 智能助手 =====
@app.route("/api/assistant", methods=["POST"])
def assistant():
    data = request.json or {}
    question = data.get("question", "").strip()
    history = data.get("history", [])

    if not question:
        return jsonify({"error": "问题为空"}), 400

    system_prompt = """你是"Offer 捕手"的智能客服助手。Offer 捕手是一个 AI 驱动的学生求职匹配工具。

功能介绍：
1. 批量匹配：上传简历 + 添加 JD → AI 多维度打分排序（经验40% + 学科25% + 技能15% + 公司10% + 适配10%），权重可调
2. 深度诊断：针对单个心仪岗位，逐项对比简历（学历/专业/经验/技能）→ 输出 pass/warn/fail + 具体修改建议
3. 简历优化（成长路线图）：AI 分析目标岗位的共性要求 → 给出短期/中期/长期提升建议
4. 投递管理：记录投递进度 → 看板视图管理（待投/已投/面试/offer）
5. 策略建议：AI 生成投递顺序 + 风险提示 + 周计划
6. 浏览器插件：在 BOSS 直聘一键抓取 JD，支持单张和批量抓取

使用流程：上传简历 → 添加 JD → 开始匹配 → 针对心仪岗位深度诊断 → 生成成长路线

常见问题：
- 如何添加 JD：在文本框粘贴岗位描述，点"+ 添加 JD"按钮，支持批量添加多条
- 简历怎么上传：可以粘贴文本，也可以拖拽上传 .pdf/.docx/.txt 文件
- 权重怎么调：在「权重设置」标签页拖滑块，不同行业可设置不同侧重
- 插件怎么装：点页面顶部"下载浏览器插件"，解压后在 chrome://extensions 加载
- 没解析出简历信息：确认简历包含教育背景、实习经历、技能等完整信息

请用友好的语气回答用户问题，回答要简洁（控制在 150 字以内），引导用户正确使用产品。"""

    messages = [{"role": "system", "content": system_prompt}]
    for h in history[-6:]:  # 保留最近 6 条对话
        messages.append({"role": h.get("role", "user"), "content": h.get("content", "")})
    messages.append({"role": "user", "content": question})

    try:
        eng = require_engine()
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    try:
        result = eng.call(json.dumps(messages, ensure_ascii=False), 500)
        answer = result.strip()
        if len(answer) > 300:
            answer = answer[:300] + "..."
        return jsonify({"ok": True, "answer": answer})
    except Exception as e:
        return jsonify({"error": f"AI 助手暂不可用: {str(e)}"}), 500


# ===== JD =====
@app.route("/api/jd/add", methods=["POST"])
def add_jd():
    try: eng = require_engine()
    except ValueError as e: return jsonify({"error": str(e)}), 400
    data = request.json or {}
    text = data.get("text", "").strip()
    if not text:
        return jsonify({"error": "JD为空"}), 400
    if len(text) < 20:
        return jsonify({"error": "JD内容太短（至少20字）"}), 400
    jd = eng.parse_jd(text)
    jd["raw"] = text
    state["jds"].append(jd)
    return jsonify({"ok": True, "jd": jd, "total": len(state["jds"]), "parsed": bool(jd.get("title") and jd["title"] != "未知岗位")})

@app.route("/api/jd/clear", methods=["POST"])
def clear_jds():
    state["jds"] = []
    return jsonify({"ok": True})

# ===== 批量匹配 =====
@app.route("/api/match", methods=["POST"])
def batch_match():
    try: eng = require_engine()
    except ValueError as e: return jsonify({"error": str(e)}), 400
    if not state["resume"]:
        return jsonify({"error": "请先上传简历"}), 400
    if not state["jds"]:
        return jsonify({"error": "请先添加JD"}), 400
    results = eng.batch_match(state["resume"], state["jds"])
    preferences = eng.analyze_preferences(state["resume"], state["jds"])
    state["last_preferences"] = preferences
    # 公司分析（取前10个）
    companies = {}
    for jd in state["jds"][:10]:
        name = jd.get("company", "")
        if name and name not in companies:
            companies[name] = eng.analyze_company(name, jd.get("industry", ""))
    return jsonify({"ok": True, "results": results, "preferences": preferences, "companies": companies, "total": len(results)})

# ===== 深度诊断 =====
@app.route("/api/diagnose", methods=["POST"])
def diagnose():
    try: eng = require_engine()
    except ValueError as e: return jsonify({"error": str(e)}), 400
    data = request.json or {}
    jd_text = data.get("jd_text", "").strip()
    if not state["resume"] or not jd_text:
        return jsonify({"error": "需要简历和JD"}), 400
    jd = eng.parse_jd(jd_text)
    result = eng.deep_diagnose(state["resume"], jd)
    return jsonify({"ok": True, "diagnosis": result})

# ===== 简历优化 =====
@app.route("/api/optimize", methods=["POST"])
def optimize():
    try: eng = require_engine()
    except ValueError as e: return jsonify({"error": str(e)}), 400
    if not state["resume"]:
        return jsonify({"error": "请先上传简历"}), 400
    target_jds = state["jds"][:5]
    if not target_jds:
        return jsonify({"error": "请先添加目标JD"}), 400
    result = eng.optimize_resume(state["resume"], target_jds)
    return jsonify({"ok": True, "optimization": result})

# ===== 投递管理 =====
@app.route("/api/track", methods=["GET"])
def get_tracks():
    tracks = state.get("tracks", {})
    return jsonify({"tracks": tracks, "jds": [{"idx": i, "title": j.get("title",""), "company": j.get("company",""), "score": j.get("score", 0)} for i, j in enumerate(state["jds"])]})

@app.route("/api/track", methods=["POST"])
def update_track():
    data = request.json or {}
    idx = data.get("index", -1)
    status = data.get("status", "pending")
    if idx < 0 or idx >= len(state["jds"]):
        return jsonify({"error": "无效的岗位索引"}), 400
    if "tracks" not in state:
        state["tracks"] = {}
    state["tracks"][str(idx)] = status
    return jsonify({"ok": True})

# ===== 投递策略 =====
@app.route("/api/strategy", methods=["POST"])
def strategy():
    try: eng = require_engine()
    except ValueError as e: return jsonify({"error": str(e)}), 400
    if not state["resume"] or not state["jds"]:
        return jsonify({"error": "请先上传简历和JD"}), 400
    jds_with_idx = [{"idx": i, **j} for i, j in enumerate(state["jds"])]
    tracks = state.get("tracks", {})
    text = json.dumps({
        "resume_summary": json.dumps(state["resume"], ensure_ascii=False)[:500],
        "jds": [{"idx": j["idx"], "title": j.get("title",""), "company": j.get("company","")} for j in jds_with_idx[:10]],
        "applied": {k: v for k, v in tracks.items()},
        "preferences": state.get("last_preferences", {})
    }, ensure_ascii=False)
    result = eng.call(
        f"你是投递策略顾问。分析候选人情况和岗位列表，给出投递策略:\n{text}\n\n"
        "输出JSON:\n"
        '{"priority_order":"投递顺序建议(先投哪些,为什么)",'
        '"batch_strategy":"分批发还是集中投,每批投哪些",'
        '"risk_alert":"需要注意的风险(如某公司偏好院校,你的简历可能吃亏)",'
        '"quick_wins":"建议优先投的2-3个最容易拿到面试的岗位",'
        '"long_shots":"值得冲但概率低的岗位",'
        '"weekly_plan":"一周投递计划建议"}',
        2000
    )
    m = re.search(r'\{[\s\S]*\}', result)
    return jsonify({"ok": True, "strategy": json.loads(m.group(0)) if m else {}})

# ===== 权重 =====
@app.route("/api/weight", methods=["GET", "POST"])
def weight():
    if request.method == "POST":
        data = request.json or {}
        state["weight"] = data
    return jsonify(state["weight"])

# ===== BOSS 插件接收 =====
@app.route("/api/jd/from_plugin", methods=["POST"])
def from_plugin():
    try: eng = require_engine()
    except ValueError as e: return jsonify({"error": str(e)}), 400
    data = request.json or {}
    jd_text = data.get("jd_text", "").strip()
    if not jd_text:
        return jsonify({"error": "JD为空"}), 400
    jd = eng.parse_jd(jd_text)
    jd["raw"] = jd_text
    jd["source"] = "BOSS直聘"
    state["jds"].append(jd)
    return jsonify({"ok": True, "jd": jd, "total": len(state["jds"])})

# ===== 视觉识别 JD（通义千问 VL — 支持多张截图拼接） =====
@app.route("/api/jd/vision", methods=["POST"])
def jd_vision():
    data = request.json or {}
    # 支持单张 image 或多张 images 数组
    images = data.get("images", [])
    if not images:
        # 兼容旧格式
        img = data.get("image", "")
        if img:
            images = [img]
    if not images:
        return jsonify({"error": "截图数据为空"}), 400

    api_key = vision_api_key()
    if not api_key:
        return jsonify({"error": "请先在网页 API 配置面板填写视觉 API Key"}), 400

    # 构建 content 数组：多张图片 + 一段提示
    content_parts = []
    for idx, img in enumerate(images):
        # 确保是完整的 data URL
        if not img.startswith("data:"):
            img = f"data:image/jpeg;base64,{img}"
        content_parts.append({
            "type": "image_url",
            "image_url": {"url": img},
        })

    # 拼接提示
    img_count = len(images)
    if img_count == 1:
        instruction = (
            "这是一个招聘网站的岗位详情截图。请从截图中提取以下信息，"
            "返回严格JSON格式（只返回JSON，不要任何解释）：\n"
            '{"title":"岗位名称","company":"公司名称",'
            '"salary":"薪资范围（如截图中有的话）",'
            '"location":"工作地点（如截图中有的话）",'
            '"jd":"岗位职责和任职要求的完整原文，保持原文格式"}\n\n'
            "注意：\n"
            "1. 只提取截图中的JD相关内容，忽略页面UI元素、导航栏、推荐列表\n"
            "2. jd字段要包含完整的岗位职责和任职要求\n"
            "3. 直接返回JSON，不要用markdown代码块包裹"
        )
    else:
        instruction = (
            f"这是{img_count}张连续的招聘网站岗位详情截图（从上到下依次排列，相邻图片有少量重叠）。"
            "请将它们拼接起来，提取完整的岗位信息。"
            "返回严格JSON格式（只返回JSON，不要任何解释）：\n"
            '{"title":"岗位名称","company":"公司名称",'
            '"salary":"薪资范围（如截图中有的话）",'
            '"location":"工作地点（如截图中有的话）",'
            '"jd":"岗位职责和任职要求的完整原文（合并所有截图，保持原文格式）"}\n\n'
            "注意：\n"
            "1. 只提取截图中的JD相关内容，忽略页面UI元素、导航栏、推荐列表\n"
            "2. jd字段要包含所有截图中出现的完整岗位职责和任职要求，不要遗漏\n"
            "3. 相邻截图的重叠部分只保留一次\n"
            "4. 直接返回JSON，不要用markdown代码块包裹"
        )
    content_parts.append({"type": "text", "text": instruction})

    try:
        resp = requests.post(
            vision_api_url(),
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
            json={
                "model": vision_model(),
                "messages": [{
                    "role": "user",
                    "content": content_parts,
                }],
            },
            timeout=60,
        )

        if resp.status_code != 200:
            return jsonify({"error": f"VL API 返回 {resp.status_code}: {resp.text[:200]}"}), 500

        result = resp.json()
        content = result.get("choices", [{}])[0].get("message", {}).get("content", "")

        # 从 VL 回复中提取 JSON
        m = re.search(r"\{[\s\S]*\}", content)
        if not m:
            return jsonify({"error": f"VL 返回格式异常: {content[:300]}"}), 500

        jd_raw = json.loads(m.group(0))
        title = jd_raw.get("title", "")
        company = jd_raw.get("company", "")
        jd_text = jd_raw.get("jd", "")
        salary = jd_raw.get("salary", "")
        location = jd_raw.get("location", "")

        if not title or not jd_text or len(jd_text) < 20:
            return jsonify({"error": f"VL 提取不完整: title={title}, jd_len={len(jd_text)}"}), 500

        # 构造完整文本发给 DeepSeek 做结构化解析
        full_text = f"【岗位名称】{title}\n【公司】{company}\n【薪资】{salary}\n【地点】{location}\n【岗位JD】\n{jd_text}"
        eng = get_engine()
        parsed = eng.parse_jd(full_text) if eng else {"title": title, "company": company}
        parsed["raw"] = full_text
        parsed["source"] = "BOSS直聘(VL)"

        # 存入内存
        state["jds"].append(parsed)

        return jsonify({
            "ok": True,
            "jd": {
                "title": title,
                "company": company,
                "salary": salary,
                "location": location,
                "jd": jd_text,
            },
            "parsed": parsed,
            "total": len(state["jds"]),
        })

    except requests.exceptions.Timeout:
        return jsonify({"error": "VL API 超时"}), 500
    except json.JSONDecodeError:
        return jsonify({"error": f"VL 返回非JSON: {content[:300] if 'content' in dir() else 'N/A'}"}), 500
    except Exception as e:
        return jsonify({"error": f"VL 调用异常: {str(e)}"}), 500


# ===== 浏览器插件下载（自动注入服务器地址） =====
@app.route("/extension/download")
def download_extension():
    import zipfile, io

    ext_dir = os.path.join(os.path.dirname(__file__), "extension")
    server_url = request.host_url.rstrip("/")

    # 读 popup.html，把 localhost 替换为实际服务器地址
    popup_path = os.path.join(ext_dir, "popup.html")
    with open(popup_path, "r", encoding="utf-8") as f:
        popup_html = f.read()
    popup_html = popup_html.replace("http://localhost:5000", server_url)

    # 内存中创建 zip
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for root, dirs, files in os.walk(ext_dir):
            for fname in files:
                filepath = os.path.join(root, fname)
                arcname = os.path.relpath(filepath, ext_dir)
                if fname == "popup.html":
                    zf.writestr(arcname, popup_html)
                else:
                    zf.write(filepath, arcname)

    buf.seek(0)
    return send_file(
        buf,
        mimetype="application/zip",
        as_attachment=True,
        download_name="offer-hunter-extension.zip",
    )


# ===== 辅助: txt/docx 文字渲染为图片（供 VL 清洗） =====
def text_to_images(text: str, max_lines: int = 60) -> list:
    """用 PIL 把文本渲染为白底黑字图片"""
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        return []
    lines = text.split("\n")
    images = []
    for chunk_start in range(0, len(lines), max_lines):
        chunk = lines[chunk_start:chunk_start + max_lines]
        w, h = 800, len(chunk) * 22 + 40
        img = Image.new("RGB", (w, h), "white")
        draw = ImageDraw.Draw(img)
        try:
            font = ImageFont.truetype("C:/Windows/Fonts/simhei.ttf", 14)
        except Exception:
            font = ImageFont.load_default()
        for i, line in enumerate(chunk):
            draw.text((15, 20 + i * 22), line, fill="black", font=font)
        import io, base64
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        images.append(base64.b64encode(buf.getvalue()).decode())
    return images


# ===== 辅助: 调 DeepSeek 文本解析简历（主力） =====
def call_ds_for_resume(text: str) -> dict:
    ds_key = get_ds_key()
    if not ds_key:
        return {}
    prompt = f"""从以下简历文本中提取结构化信息。严格返回JSON（只返回JSON，不要markdown包裹，不要省略字段）：
{{"name":"","phone":"","email":"",
 "education":[{{"school":"","degree":"","major":"","start":"","end":""}}],
 "experience":[{{"company":"","title":"","start":"","end":"","description":""}}],
 "projects":[{{"name":"","role":"","description":""}}],
 "campus":[{{"org":"","title":"","description":""}}],
 "skills":[],"awards":[],"certificates":[]}}

简历文本：
{text[:4000]}
"""
    resp = requests.post(
        api_config["text_url"],
        headers={"Authorization": f"Bearer {ds_key}", "Content-Type": "application/json"},
        json={"model": api_config["text_model"], "temperature": 0, "max_tokens": 3000,
              "messages": [{"role": "user", "content": prompt}]},
        timeout=30,
    )
    if resp.status_code != 200:
        return {}
    content = resp.json().get("choices", [{}])[0].get("message", {}).get("content", "")
    # 多策略 JSON 提取
    parsed = None
    # 1. 直接 parse
    try:
        parsed = json.loads(content.strip())
    except json.JSONDecodeError:
        pass
    # 2. 取第一个 { 到最后一个 }
    if not parsed:
        m = re.search(r"\{[\s\S]*\}", content)
        if m:
            try:
                parsed = json.loads(m.group(0))
            except json.JSONDecodeError:
                pass
    # 3. 去掉 markdown 代码块再试
    if not parsed:
        clean = re.sub(r"```(?:json)?\s*|\s*```", "", content).strip()
        try:
            parsed = json.loads(clean)
        except json.JSONDecodeError:
            m = re.search(r"\{[\s\S]*\}", clean)
            if m:
                try:
                    parsed = json.loads(m.group(0))
                except json.JSONDecodeError:
                    pass
    return parsed or {}


# ===== 辅助: 调 VL 解析简历（兜底：扫描版 PDF / 图片） =====
def call_vl_for_resume(images_b64: list) -> dict:
    api_key = vision_api_key()
    if not api_key:
        return {}
    content_parts = []
    for img in images_b64:
        content_parts.append({
            "type": "image_url",
            "image_url": {"url": f"data:image/png;base64,{img}"},
        })
    content_parts.append({
        "type": "text",
        "text": (
            "这是一份简历。请提取以下结构化信息，返回JSON（只返回JSON，不要加任何解释）：\n"
            '{"name":"姓名","phone":"电话","email":"邮箱",'
            '"education":[{"school":"学校","degree":"学位","major":"专业","start":"开始时间","end":"结束时间"}],'
            '"experience":[{"company":"公司","title":"职位","start":"开始时间","end":"结束时间","description":"工作描述"}],'
            '"projects":[{"name":"项目名称","role":"角色","description":"项目描述"}],'
            '"campus":[{"org":"组织/社团","title":"职务","description":"活动描述"}],'
            '"skills":["技能1"],"awards":["奖项1"],"certificates":["证书1"]}'
        ),
    })
    resp = requests.post(
        api_config["vision_url"],
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        json={"model": api_config["vision_model"], "messages": [{"role": "user", "content": content_parts}]},
        timeout=60,
    )
    if resp.status_code != 200:
        return {}
    content = resp.json().get("choices", [{}])[0].get("message", {}).get("content", "")
    m = re.search(r"\{[\s\S]*\}", content)
    if m:
        return json.loads(m.group(0))
    return {}


if __name__ == "__main__":
    print("=" * 40)
    print("  Offer 捕手 v1.0")
    print("  http://localhost:5000")
    print("=" * 40)
    app.run(host="0.0.0.0", port=5000, debug=False)
